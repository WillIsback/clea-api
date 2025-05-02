from __future__ import annotations

from datetime import date

from docx import Document as DocxDocument

from ..base import (  # helpers communs
    BaseExtractor,
    build_document_with_chunks,
    DocumentWithChunks,
)


class DocxExtractor(BaseExtractor):
    """
    Convertit un fichier **.docx** en un unique payload
    `DocumentWithChunks` prêt à être POSTé sur `/database/documents`.
    """

    def __init__(self, file_path: str) -> None:
        """Initialise l'extracteur DOCX avec les métadonnées par défaut.

        Args:
            file_path (str): Chemin vers le fichier DOCX.

        Raises:
            ValueError: Si le fichier DOCX ne peut pas être ouvert.
        """
        super().__init__(file_path)
        try:
            self._docx = DocxDocument(file_path)
        except Exception as e:
            raise ValueError(f"Erreur lors de l'ouverture du fichier DOCX : {e}")

        self.default_meta = dict(
            title=str(self._guess_title()),  # Convertit le titre en chaîne
            theme="Générique",  # Déjà une chaîne
            document_type="DOCX",  # Déjà une chaîne
            publish_date="",
        )

    def _guess_title(self) -> str:
        """Devine le titre du document à partir des propriétés ou du nom de fichier.

        Returns:
            str: Titre du document.
        """
        try:
            props = self._docx.core_properties
            return props.title or self.file_path.stem
        except Exception:  # pragma: no cover
            return self.file_path.stem

    def extract_one(self, max_length: int = 1_000) -> DocumentWithChunks:
        """
        Extrait un seul objet `DocumentWithChunks` à partir du fichier DOCX.

        Args:
            max_length (int): Longueur cible pour les segments finaux.

        Returns:
            DocumentWithChunks: Objet contenant les chunks et métadonnées.
        """
        if not self._docx.paragraphs:  # DOCX vide
            raise ValueError("Le document DOCX est vide")

        full_text = "\n".join(p.text for p in self._docx.paragraphs).strip()

        return build_document_with_chunks(
            title=self.default_meta["title"],
            theme=self.default_meta["theme"],
            document_type=self.default_meta["document_type"],
            publish_date=date.today(),
            max_length=max_length,
            full_text=full_text,
        )
