from pathlib import Path
from docx import Document
from fpdf import FPDF
import json

"""
Generate demonstration files in various formats (TXT, JSON, DOCX, PDF, HTML).

:param output_dir: Directory where the demo files will be created.
:type output_dir: Path
:raises IOError: If there is an issue writing to the output directory.
:return: None
:rtype: None
"""

# Répertoire de sortie
output_dir = Path(".")
output_dir.mkdir(parents=True, exist_ok=True)

# Contenu de base pour chaque fichier (un peu plus de 1000 caractères)
sample_text = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 20
    + "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. " * 10
)

# 1. TXT
txt_path = output_dir / "demo.txt"
txt_path.write_text(sample_text, encoding="utf-8")

# 2. JSON
json_path = output_dir / "demo.json"
json_data = [
    {
        "title": "Document JSON 1",
        "content": sample_text,
        "theme": "Démo",
        "document_type": "JSON",
        "publish_date": "2024-01-01",
    },
    {
        "title": "Document JSON 2",
        "content": sample_text,
        "theme": "Démo",
        "document_type": "JSON",
        "publish_date": "2024-01-02",
    },
]
json_path.write_text(json.dumps(json_data, indent=2), encoding="utf-8")

# 3. DOCX
docx_path = f"{output_dir}/demo.docx"
doc = Document()
doc.add_heading("Document DOCX de démonstration", 0)
doc.add_paragraph(sample_text)
doc.save(docx_path)

# 4. PDF
pdf_path = output_dir / "demo.pdf"
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
for line in sample_text.split(". "):
    pdf.multi_cell(0, 10, line)
pdf.output(str(pdf_path))

# 5. HTML
html_path = output_dir / "demo.html"
html_content = f"""
<!DOCTYPE html>
<html lang="fr">
<head><meta charset="UTF-8"><title>Document HTML</title></head>
<body>
<h1>Document HTML de démonstration</h1>
<p>{sample_text}</p>
</body>
</html>
"""
html_path.write_text(html_content, encoding="utf-8")
